"""CVBuilder handling component

This component is used to generate a CV document from a given yaml description"""

import io
import logging

import fontawesome as fa
from docx import Document
from docx2pdf import convert
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from nvp.nvp_component import NVPComponent
from nvp.nvp_context import NVPContext

logger = logging.getLogger(__name__)


class CVBuilder(NVPComponent):
    """CVBuilder component class"""

    def __init__(self, ctx: NVPContext):
        """Component constructor"""
        NVPComponent.__init__(self, ctx)
        self.document = None
        self.desc = None

    def process_cmd_path(self, cmd):
        """Re-implementation of process_cmd_path"""

        if cmd == "build":
            file = self.get_param("input_file")

            # Read the configuration:
            cfg = self.read_yaml(file)

            return self.build(cfg)

        return False

    def convert_icon_to_image(self, icon_name, size=32, color="black"):
        """Convert a fontawesome icon to an image"""
        # Specify the Font Awesome font path
        # font_path = fa.font_path()

        # Set the font file path
        # font_file = "fonts/fa-regular-400.ttf"
        font_file = "fonts/fa-solid-900.ttf"

        # Set the icon and size
        # logger.info("Available icons: %s", fa.icons.keys())
        icon = fa.icons[icon_name]

        # Create a blank image with an alpha channel
        image = Image.new("RGBA", (size, size), (0, 0, 0, 0))

        # Create a drawing context
        draw = ImageDraw.Draw(image)

        # Load the Font Awesome font
        font = ImageFont.truetype(font_file, size - 6)

        # Calculate the text size and position
        text_width, text_height = draw.textsize(icon, font=font)
        text_x = (size - text_width) // 2
        text_y = (size - text_height) // 2

        # Draw the icon on the image
        draw.text((text_x, text_y), icon, font=font, fill=color)

        image.save("icon_image.png", "PNG")

        return image

    def add_image(self, element, img, width):
        """Add a PIL image to an element"""
        image_stream = io.BytesIO()
        img.save(image_stream, format="PNG")
        return element.add_picture(image_stream, width=width)

    def build(self, desc):
        """This function is used build the CV from the given description"""
        filename = desc["filename"]
        docx_file = filename + ".docx"
        pdf_file = filename + ".pdf"

        # Create a new document
        self.document = Document()
        self.desc = desc

        # self.build_content1(document)
        # self.build_content2(self.document)
        self.build_content()

        # Save the CV as a DOCX file
        self.document.save(docx_file)
        self.document = None

        # Convert the DOCX file to PDF
        convert(docx_file, pdf_file)

        logger.info("Done writing %s", docx_file)
        return True

    def build_content(self):
        """Build the final content for the CV"""

        # Remove the margins:
        section = self.document.sections[-1]
        section.left_margin = 0
        section.right_margin = 0
        section.top_margin = 0
        section.bottom_margin = 0

        # Define the styles
        self.define_styles()

        # Prepare the 2 columns:
        table = self.document.add_table(rows=1, cols=2)
        table.style = "Table Grid"  # Apply table grid style
        table.autofit = False
        table.allow_autofit = False

        # Build the side column:
        self.build_side_column(table.cell(0, 0))

        # Build the main column:
        self.build_main_column(table.cell(0, 1))

        # Set the column widths for each cell in each column
        column_widths = (self.document.sections[0].page_width * 0.2, self.document.sections[0].page_width * 0.8)
        logger.info("Setting column widths to %f, %f", column_widths[0], column_widths[1])
        self.set_col_widths(table, column_widths)
        self.remove_cell_borders(table)

    def add_paragraph_style(self, style_name, font_name, font_size, color):
        """Add a paragraph style"""
        style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(color[0], color[1], color[2])
        return style

    # Not working:
    # def set_style_font_file(self, style, font_file):
    #     """Assign the font file to a style"""
    #     style_element = style.element
    #     rpr = style_element.get_or_add_rPr()

    #     r_fonts = rpr.first_child_found_in("w:rFonts")

    #     if r_fonts is None:
    #         logger.info("Adding custom font to style...")
    #         r_fonts = OxmlElement("w:rFonts")
    #         rpr.append(r_fonts)

    #     r_fonts.set(qn("w:ascii"), font_file)
    #     r_fonts.set(qn("w:hAnsi"), font_file)
    #     r_fonts.set(qn("w:eastAsia"), font_file)

    def define_styles(self):
        """Define the styles we are going to use in the document"""
        # self.add_paragraph_style("MainTitle", "Tahoma", 32, (0, 0, 0))
        # self.add_paragraph_style("MainTitle", "Verdana", 40, (0, 0, 0))
        # self.add_paragraph_style("MainTitle", "Helvetica", 40, (0, 0, 0))
        # self.add_paragraph_style("MainTitle", "Roboto", 40, (0, 0, 0))
        style = self.add_paragraph_style("MainTitle", "Calibri", 40, (0, 0, 0))
        style = self.add_paragraph_style("Qualifications", "Tahoma", 9, (0, 110, 184))
        style.font.small_caps = True
        # style = self.add_paragraph_style("Address", "Tahoma", 9, (153, 153, 153))
        # self.set_style_font_file(style, "Roboto-Regular.ttf")
        # self.embed_font("Stroke.ttf")
        style = self.add_paragraph_style("Address", "Tahoma", 9, (153, 153, 153))
        # self.set_style_font_file(style, "Stroke.ttf")

        style = self.add_paragraph_style("PhoneEmail", "Tahoma", 8, (51, 51, 51))

    def build_main_column(self, cell):
        """Build the content of the main column"""

        first_name = self.desc["first_name"]
        last_name = self.desc["last_name"]

        pgraph = cell.add_paragraph(first_name + " ", style="MainTitle")
        pgraph.add_run(last_name).bold = True

        content = " · ".join(self.desc["qualifications"])
        pgraph = cell.add_paragraph(content, style="Qualifications")

        pgraph = cell.add_paragraph(self.desc["address"], style="Address")

        # table = cell.add_table(rows=1, cols=2)
        # # table.style = "Table Grid"
        # table.allow_autofit = True
        # table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # # Get the first row of the table
        # cell = table.cell(0, 0)

        # add the phone/email line:
        phone_str = self.desc["phone"]
        email_str = self.desc["email"]
        phone_img = self.convert_icon_to_image("phone", 256)
        email_img = self.convert_icon_to_image("envelope", 256)
        phone_img.save("phone.png", "PNG")
        email_img.save("email.png", "PNG")

        width = Inches(0.28)
        height = width * phone_img.height / phone_img.width

        # # Add content to the cells
        # #

        # # run = cell1.add_paragraph(style="PhoneEmail").add_run()
        # pgraph = cell.paragraphs[0]
        # # pgraph.style = "PhoneEmail"
        # run = pgraph.add_run()
        # run.add_picture("phone.png", width=width, height=height)
        # cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # # ww1 = run.width
        # # hh1 = run.height
        # # logger.info("Run1 size: %dx%d", ww1, hh1)

        # cell = table.cell(0, 1)
        # # run = cell2.add_paragraph(style="PhoneEmail").add_run()
        # pgraph = cell.paragraphs[0]
        # # pgraph.style = "PhoneEmail"
        # run = pgraph.add_run()
        # run.add_text("My phone number")
        # cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # # ww2 = run.width
        # # hh2 = run.height
        # # logger.info("Run2 size: %dx%d", ww2, hh2)

        # table._tblPr.xpath("./w:tblW")[0].attrib[
        #     "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
        # ] = "auto"
        # for row in table.rows:
        #     for cell in row.cells:
        #         cell.width = 0
        #         tc = cell._tc
        #         tcPr = tc.get_or_add_tcPr()
        #         tcW = tcPr.get_or_add_tcW()
        #         tcW.type = "auto"
        #         tcW.w = 0

        # # Set vertical alignment of cells to center
        # # cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # table.autofit = True
        pgraph = cell.add_paragraph(style="PhoneEmail")
        run = pgraph.add_run()
        pic = run.add_picture("phone.png", width=width, height=height)
        pic.left_margin = 0
        pic.right_margin = 0
        pic.space_before = Pt(0)
        pic.space_after = Pt(0)

        # image_run.bottom_margin = -Inches(0.05)
        #
        # run = pgraph.add_run("My phone number")

        # # Adjust the vertical alignment of the runs
        # image_run.vertical_alignment = "center"
        # text_run.vertical_alignment = "center"

        # self.add_image(run, phone_img, 32)
        # pgraph.add_run(phone_str + " | ")
        run.add_text(" " + phone_str + " | ")
        # run.add_picture("email.png", width=32)
        pic = run.add_picture("email.png", width=width, height=height)
        pic.left_margin = 0
        pic.right_margin = 0
        pic.space_before = Pt(0)
        pic.space_after = Pt(0)

        # # self.add_image(run, email_img, 32)
        run.add_text(" " + email_str)
        # pgraph.add_run().add_picture("email.png")
        # pgraph.add_run(" | ")
        # pgraph.add_run(email_str)

    def build_side_column(self, cell):
        """Build the content of theside column"""

        cell.add_paragraph("This is the side colum")

    def build_content1(self, document):
        "Test build content 1"

        # Add a heading with the full name
        full_name = "John Doe"
        document.add_heading(full_name, level=1)

        # Add contact information
        email = "johndoe@example.com"
        phone = "+1 234 567 890"
        address = "123 Main Street, City, Country"
        self.add_contact_information(document, email, phone, address)

        # Add a summary section
        summary = "A brief summary of your skills and experience."
        self.add_summary_section(document, summary)

    def build_content2(self, document):
        "Test build content 2"

        section = document.sections[-1]
        # section.start_type = 1  # Continuous section break

        # # Define column widths
        # column_widths = (section.page_width * 0.2, section.page_width * 0.8)
        # section.column_widths = column_widths

        # Adjust margins to remove spacing between columns
        section.left_margin = 0
        section.right_margin = 0
        section.top_margin = 0
        section.bottom_margin = 0

        # Add a table with one row and two columns
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"  # Apply table grid style
        table.autofit = False
        table.allow_autofit = False
        table.left_margin = 0
        table.right_margin = 0
        table.top_margin = 0
        table.bottom_margin = 0

        # self.embed_font("fonts/Coval-Regular.otf")

        # Add content to the document
        self.add_content(table)

        # Set the column widths for each cell in each column
        column_widths = (document.sections[0].page_width * 0.2, document.sections[0].page_width * 0.8)
        logger.info("Setting column widths to %f, %f", column_widths[0], column_widths[1])
        self.set_col_widths(table, column_widths)
        self.remove_cell_borders(table)

        # table.columns[0].width = int(column_widths[0])
        # table.columns[1].width = int(column_widths[1])

    def create_custom_style(self, style_name, font_size, font_color):
        """Create a new custom style element"""
        style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])
        return style

    def set_col_widths(self, table, widths):
        """Apply the column widths"""
        for idx, col in enumerate(table.columns):
            col.width = int(widths[idx])

        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = int(width)

    def remove_cell_borders(self, table):
        """Remove the cell borders"""

        data = {"sz": 0, "val": "none", "color": "#FF0000", "space": "0", "shadow": "false"}
        for row in table.rows:
            for cell in row.cells:
                self.set_cell_borders(cell, start=data, end=data, top=data, bottom=data)

    def set_cell_borders(self, cell, **kwargs):
        """Remove the cell borders"""
        # cf. https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
        # cf. http://officeopenxml.com/WPtableBorders.php

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        # if tcBorders is not None:
        #     logger.info("Removing borders on cell...")
        #     tcPr.remove(tcBorders)

        if tcBorders is None:
            logger.info("Adding borders on cell...")
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = "w:{}".format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn("w:{}".format(key)), str(edge_data[key]))

    def add_content(self, table):
        """Add content to our table"""

        # Add text content to the table cells
        table.cell(0, 0).paragraphs[0].add_run("Column 1").bold = True
        table.cell(0, 1).paragraphs[0].add_run("Column 2").bold = True

        # Define custom style properties
        style_name = "MyStyle"
        style_font_size = 14
        style_font_color = (255, 0, 0)  # Red color

        # Create a custom style element
        self.create_custom_style(style_name, style_font_size, style_font_color)

        # Apply the custom style to a paragraph
        paragraph = table.cell(0, 1).add_paragraph("This is a paragraph with custom style")
        paragraph.style = "MyStyle"

        # Add more content to the table cells as desired
        # You can modify or add paragraphs, runs, etc.

    def add_contact_information(self, document, email, phone, address):
        """Add contact information"""
        document.add_paragraph("Contact Information", style="Heading2")

        # Add email
        document.add_paragraph(f"Email: {email}", style="BodyText")

        # Add phone
        document.add_paragraph(f"Phone: {phone}", style="BodyText")

        # Add address
        document.add_paragraph(f"Address: {address}", style="BodyText")

        document.add_paragraph()  # Add an empty paragraph for spacing

    def add_summary_section(self, document, summary):
        """Add summary section"""
        document.add_paragraph("Summary", style="Heading2")
        document.add_paragraph(summary, style="BodyText")


if __name__ == "__main__":
    # Create the context:
    context = NVPContext()

    # Add our component:
    comp = context.register_component("CVBuilder", CVBuilder(context))

    psr = context.build_parser("build")
    psr.add_str("-i", "--input", dest="input_file")("Input CV yaml config file to use")

    comp.run()
